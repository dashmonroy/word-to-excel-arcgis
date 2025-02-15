import os
import re
import pandas as pd
from docx import Document
from datetime import datetime
from multiprocessing import Pool

# 📌 Rutas de los archivos
script_dir = os.path.dirname(os.path.abspath(__file__))
word_folder = os.path.join(script_dir, "Archivos")
timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
output_excel = os.path.join(script_dir, f"hallazgos_extraidos_{timestamp}.xlsx")

# 📌 Mapeo de nombres de meses en español a números
meses = {
    "enero": "01", "febrero": "02", "marzo": "03", "abril": "04", "mayo": "05", "junio": "06",
    "julio": "07", "agosto": "08", "septiembre": "09", "octubre": "10", "noviembre": "11", "diciembre": "12"
}

# 📌 Función para convertir fecha al formato requerido
def convertir_fecha(fecha_texto):
    try:
        fecha_texto = fecha_texto.strip()
        patrones = [
            r"(\d{1,2}) de (\w+) de (\d{4}) (\d{1,2}:\d{2})",
            r"(\d{1,2})/(\d{1,2})/(\d{4}) (\d{1,2}:\d{2})"
        ]
        for patron in patrones:
            match = re.search(patron, fecha_texto)
            if match:
                grupos = match.groups()
                if len(grupos) == 4:
                    dia, mes_texto, anio, hora = grupos
                    mes = meses.get(mes_texto.lower(), "00")
                else:
                    dia, mes, anio, hora = grupos
                return f"{dia.zfill(2)}/{mes}/{anio} {hora}"
    except:
        pass
    return "Formato Incorrecto"

# 📌 Función para extraer texto de documentos Word
def extract_text_from_docx(docx_path):
    doc = Document(docx_path)
    text = "\n".join([para.text.strip() for para in doc.paragraphs if para.text.strip()])
    for table in doc.tables:
        for row in table.rows:
            row_text = " | ".join([cell.text.strip() for cell in row.cells])
            text += "\n" + row_text
    return text.strip()

# 📌 Extraer información clave con expresiones regulares
def extract_field(text, field):
    match = re.search(rf"{field}[:\s]*(.+)", text, re.IGNORECASE)
    return match.group(1).strip() if match else "No encontrado"

# 📌 Campos a extraer
fields = [
    "FECHA Y HORA", "NOMBRE QUIEN REPORTA", "ALIADO", "REGIONAL",
    "DEPARTAMENTO", "LISTA DE MUNICIPIOS", "MUNICIPIO", "BARRIO",
    "DIRECCIÓN", "PUNTO DE REFERENCIA", "VULNERABILIDADES DE INFRAESTRUCTURA",
    "ESTADO DE INFRAESTRUCTURA (TAPA/ RECAMARA ) INDIQUE EL NIVEL",
    "DUEÑO INFRAESTRUCTURA", "CODIGO PACVI", "VULNERABILIDADES RED PROPIA",
    "Asociar OT", "Prioridad de vulnerabilidad", "OBSERVACIONES SOBRE EL HALLAZGO"
]

# 📌 Procesamiento de archivos Word
def process_word(word_file):
    word_path = os.path.join(word_folder, word_file)
    text = extract_text_from_docx(word_path)
    extracted_data = {"Nombre Archivo": word_file}
    for field in fields:
        extracted_data[field] = extract_field(text, field)
    if "FECHA Y HORA" in extracted_data:
        extracted_data["FECHA Y HORA"] = convertir_fecha(extracted_data["FECHA Y HORA"])
    return extracted_data

# 📌 Procesamiento en paralelo
if __name__ == "__main__":
    word_files = [f for f in os.listdir(word_folder) if f.endswith(".docx")]
    with Pool(4) as p:
        word_data = p.map(process_word, word_files)
    df = pd.DataFrame(word_data)
    df.to_excel(output_excel, index=False)
    print(f"✅ Datos extraídos con éxito y guardados en: {output_excel}")
